
# Import the required libraries
import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
import joblib
import nltk
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
import requests
import io
import docx2pdf
import os
import win32com.client
import tempfile
import re

nltk.download('stopwords')
nltk.download('punkt')

# Load pickle files
model = joblib.load('model_MNBD.pkl')
tfidf_vectorizer = joblib.load('VECTOR.pkl')

# Data preprocessing and text cleaning
def preprocess_resume(text):
    # Convert to lowercase
    text = text.lower()
    text = re.sub('http\S+\s*', ' ', text)  # remove URLs
    text = re.sub('#\S+', '', text)  # remove hashtags
    text = re.sub(r'\d+', '', text)  # Remove numbers
    text = re.sub('@\S+', '  ', text)  # remove mentions
    text = re.sub(r'\b\d+\b', '', text)  # remove independent numbers
    text = re.sub('[%s]' % re.escape("""!"#$%&'()*+,-./:;<=>?@[\]^_`{|}~"""), ' ', text)  # remove punctuations
    text = re.sub(r'[^\x00-\x7f]', ' ', text)  # remove non-Latin characters
    text = re.sub('\s+', ' ', text)  # remove extra whitespace
    # Tokenization (split text into words)
    words = nltk.word_tokenize(text)

    # Remove stopwords
    stop_words = set(stopwords.words('english'))
    words = [word for word in words if word not in stop_words]

    # Join words back into a sentence
    processed_text = ' '.join(words)

    return processed_text


# Function to convert .doc to .docx using python-docx library
def convert_to_docx(file):
    try:
        # Save the .doc file to a temporary file
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp_file:
            tmp_filename = tmp_file.name
            tmp_file.write(file.read())

        # Define the name and path of the corresponding .docx file
        docx_filename = tmp_filename + "x"

        # Create a new instance of the Word application
        word_app = win32com.client.Dispatch("Word.Application")

        # Open the .doc file
        doc = word_app.Documents.Open(tmp_filename)

        # Save the document in .docx format
        doc.SaveAs2(docx_filename, FileFormat=16)

        # Close the document
        doc.Close()

        # Quit the Word application
        word_app.Quit()

        # Read the converted .docx file
        with open(docx_filename, "rb") as docx_file:
            doc = Document(docx_file)
            text = ' '.join([para.text for para in doc.paragraphs])

        # Remove the temporary files
        os.remove(tmp_filename)
        os.remove(docx_filename)

        return text
    except Exception as e:
        st.error(f'Error occurred while converting .doc to .docx: {str(e)}')
        return None


# Function to process the uploaded file and make predictions
def process_file(file):
    text = ''
    filename, file_extension = os.path.splitext(file.name)
    
    if file_extension == '.pdf':
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    elif file_extension == '.docx':
        doc = Document(io.BytesIO(file.read()))
        for para in doc.paragraphs:
            text += para.text
    elif file_extension == '.doc':
        try:
            text = convert_to_docx(file)
            if text is None:
                return
        except Exception as e:
            st.error(f'Error occurred while converting .doc to .docx: {str(e)}')
            return
    else:
        st.error('Unsupported file format. Please upload a PDF or DOC/DOCX file.')
        return

    preprocessed_text = preprocess_resume(text)
    vectorized_text = tfidf_vectorizer.transform([preprocessed_text])
    prediction = model.predict(vectorized_text)
    return prediction[0]


# Creating the Streamlit app and defining its behavior
def main():
    st.title('Resume Classification')
    st.write('Upload a resume in PDF or DOC/DOCX format to classify.')

    file = st.file_uploader('Upload File', type=['pdf', 'docx', 'doc'])

    if file is not None:
        prediction = process_file(file)
        st.write('Prediction:', prediction)


if __name__ == '__main__':
    main()
